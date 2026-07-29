"""Build the public ContextGuard Word progress report.

This is intentionally a reproducible builder: data and raw reviewer runs stay in
the external data home; only aggregate artifacts and small illustrative excerpts
are embedded in the report.
"""

# The report contains long Vietnamese prose and source URLs; keep those lines
# readable in the generated document instead of mechanically wrapping strings.
# ruff: noqa: E501

from __future__ import annotations

import csv
import json
import statistics
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from context_guard import ContextGuard, GuardConfig
from context_guard.adapters import RuleBasedCompressor
from context_guard.storage import get_data_home

ROOT = Path(__file__).resolve().parents[1]
DATA_HOME = get_data_home()
REPORT_DIR = ROOT / "reports"
ASSET_DIR = ROOT / ".runtime" / "report_assets"
REPORT_PATH = REPORT_DIR / "ContextGuard_Progress_Report.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF8E8"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def _font(name: str = "Calibri", size: float = 11, bold: bool | None = None, color: str | None = None):
    def apply(run: Any) -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    return apply


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size: float = 9.2) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    _set_table_geometry(table, widths)
    _set_repeat_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _shade(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        _font(size=font_size, bold=True, color=INK)(run)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=True):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(str(value))
            _font(size=font_size, color="1F2937")(run)
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, style: str | None = None) -> Any:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        _font(bold=True)(first)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        _font()(rest)
    else:
        _font()(paragraph.add_run(text))
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    _font()(paragraph.add_run(text))


def add_callout(doc: Document, label: str, text: str, fill: str = CAUTION) -> None:
    table = doc.add_table(rows=1, cols=1)
    _set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    _font(size=10.5, bold=True, color=INK)(label_run)
    _font(size=10.5, color=INK)(paragraph.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(doc: Document, text: str, level: int = 1) -> Any:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    return paragraph


def add_page_number(paragraph: Any) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    _font(size=9, color=MUTED)(run)


def _load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _load_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def make_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 28)
        font_small = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 18)
        font_title = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 32)
    except OSError:
        font = ImageFont.load_default()
        font_small = font
        font_title = font
    pipeline = ASSET_DIR / "pipeline.png"
    image = Image.new("RGB", (1600, 560), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        (60, 190, 300, 370, "Original\ncontext"),
        (360, 190, 600, 370, "Compressor\nR6-R9"),
        (660, 190, 900, 370, "Deterministic\nGuard"),
        (960, 190, 1200, 370, "Semantic\nUNCERTAIN only"),
        (1260, 190, 1540, 370, "PASS / FAIL /\nUNCERTAIN + audit")
    ]
    colors = ["DCEAF7", "E8EEF5", "FFF8E8", "F4F6F9", "DCEFE3"]
    for (x1, y1, x2, y2, label), fill in zip(boxes, colors, strict=True):
        draw.rounded_rectangle(
            (x1, y1, x2, y2), radius=22, fill=f"#{fill}", outline="#2E74B5", width=3
        )
        lines = label.split("\n")
        for index, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text(
                ((x1 + x2 - (bbox[2] - bbox[0])) / 2, y1 + 66 + index * 26),
                line,
                fill="#0B2545",
                font=font,
            )
    for start in (300, 600, 900, 1200):
        draw.line((start + 10, 280, start + 55, 280), fill="#1F4D78", width=6)
        draw.polygon(
            [(start + 55, 280), (start + 38, 270), (start + 38, 290)], fill="#1F4D78"
        )
    draw.text(
        (60, 60),
        "ContextGuard end-to-end safety pipeline",
        fill="#0B2545",
        font=font_title,
    )
    image.save(pipeline)

    metrics = list(csv.DictReader((ROOT / "artifacts/final/m13_compressor_metrics.csv").open(encoding="utf-8")))
    grouped: dict[str, list[float]] = defaultdict(list)
    for row in metrics:
        if row["mode"] == "R7_direct_guard" and row["protected"] == "True":
            grouped[f"{row['compressor']}\n{row['rate']}"] .append(float(row["safe_saving_ratio"]))
    chart = ASSET_DIR / "safe_saving.png"
    image = Image.new("RGB", (1400, 720), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (70, 35),
        "Safe effective token saving after deterministic guard",
        fill="#0B2545",
        font=font_title,
    )
    base_x, base_y, chart_w, chart_h = 100, 620, 1200, 500
    draw.line((base_x, base_y, base_x + chart_w, base_y), fill="#667085", width=2)
    draw.line((base_x, base_y, base_x, base_y - chart_h), fill="#667085", width=2)
    max_value = max(grouped.values(), default=[0.1])[0] if grouped else 0.1
    max_value = max(max_value, 0.1)
    labels = list(grouped)
    bar_w = max(20, int(chart_w / max(len(labels), 1) * 0.65))
    for idx, label in enumerate(labels):
        value = statistics.mean(grouped[label])
        height = int(value / max_value * chart_h)
        x = base_x + int((idx + 0.5) * chart_w / len(labels))
        draw.rectangle(
            (x - bar_w // 2, base_y - height, x + bar_w // 2, base_y), fill="#2E74B5"
        )
        draw.text(
            (x - bar_w // 2, base_y + 12),
            label.replace("\n", " "),
            fill="#1F2937",
            font=font_small,
        )
        draw.text(
            (x - bar_w // 2, base_y - height - 18),
            f"{value:.3f}",
            fill="#0B2545",
            font=font_small,
        )
    image.save(chart)
    return {"pipeline": pipeline, "safe_saving": chart}


def build() -> Path:
    assets = make_assets()
    dev = _load_jsonl(DATA_HOME / "normalized" / "natural_adversarial_v1_dev.jsonl")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("ContextGuard | Progress & Evidence Report")
    _font(size=9, color=MUTED)(header_run)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Public Git artifact  |  Page ")
    _font(size=9, color=MUTED)(footer_run)
    add_page_number(footer)

    # Cover / memo masthead. The static TOC is inserted before this block;
    # start the cover on a clean page after the TOC.
    doc.add_page_break()
    add_body(doc, "TECHNICAL PROGRESS REPORT", bold_prefix="TECHNICAL PROGRESS REPORT")
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("ContextGuard")
    _font(size=28, bold=True, color=INK)(title_run)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _font(size=14, color=MUTED)(subtitle.add_run("Offline-first safety validation for compressed context"))
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Repository", "https://github.com/Klein1411/ContextGuard"],
            ["Report date", "29 July 2026"],
            ["Scope", "M0-M14 expansion; M10a, M11, M12 and M13 evidence"],
            ["Decision status", "COMPLETED_WITH_LIMITATIONS; M14 gates passed"],
            ["Public rule", "Exactly three Markdown files; raw/model/reviewer data external"],
        ],
        [2000, 7360],
        font_size=9.5,
    )
    add_callout(
        doc,
        "Executive decision",
        "The goal is technically feasible and has been executed through the bounded M14 hand-off. Treat the result as a public, evidence-backed engineering milestone—not a production or paper-grade quality claim.",
        fill="EAF3F8",
    )
    doc.add_page_break()

    add_heading(doc, "1. Mục lục", 1)
    add_body(doc, "Mục lục tĩnh được chèn sau khi toàn bộ Heading 1-3 đã được tạo; các mục có liên kết nội bộ trong Word.")
    add_heading(doc, "2. Tóm tắt điều hành", 1)
    add_body(doc, "ContextGuard giữ đúng boundary đã chốt: core deterministic/offline, chỉ nhận text, kiểm tra fact và logic sau nén, rồi fallback an toàn khi FAIL hoặc UNCERTAIN. Goal mở rộng không xung đột với API V1 vì M10a chỉ đo R1-R5, còn R6-R9 được chuyển hẳn sang M13.")
    add_table(
        doc,
        ["Kết luận", "Bằng chứng", "Giới hạn"],
        [
            ["Khả thi", "M10a chạy R1-R5; M11 400 mẫu; M12 5×400 blind AI; M13 672 rows", "Pilot bounded, chưa production"],
            ["An toàn kiến trúc", "Semantic chỉ gọi trên UNCERTAIN; protected spans; fallback explicit", "NLI pair limit 512"],
            ["Có thể public", "Manifest/checksum/artifact public; raw data ngoài D:\\fact_safeguard_data", "QASPER payload unavailable"],
        ],
        [2300, 4300, 2760],
    )
    add_callout(doc, "Trạng thái", "COMPLETED_WITH_LIMITATIONS: M14 đã xác nhận test, security, git, data-home và Word artifact; các limitation được giữ nguyên trong báo cáo.")

    add_heading(doc, "3. Mục tiêu, phạm vi và các chỉnh sửa scope", 1)
    add_heading(doc, "3.1 Mục tiêu sản phẩm", 2)
    add_body(doc, "Mục tiêu cốt lõi là một lớp ContextGuard có thể đứng giữa compressor và downstream model: original context đi vào, candidate context sau nén được kiểm tra, rồi hệ thống trả PASS/FAIL/UNCERTAIN và recommended action. Core không tự gọi remote service và không phụ thuộc compressor cụ thể.")
    add_heading(doc, "3.2 Scope đã khóa để tránh xung đột", 2)
    for item in (
        "M10a chỉ đo direct deterministic, ASGI in-process, localhost HTTP, Docker HTTP và direct hybrid semantic; R6-R9 end-to-end compressor chỉ thuộc M13.",
        "M11 phân biệt natural, translated và adversarial; marker bản dịch không được coi là naturally Vietnamese.",
        "M12 dùng thuật ngữ blind multi-agent AI review/AI-adjudicated; không giả danh human/domain-expert review.",
        "M13 dùng gross saving và safe effective saving; safe saving là metric chính, break-even là mục tiêu phụ.",
        "Word report là artifact public; raw data, cache, checkpoints và reviewer runs nằm ngoài Git.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "4. Pipeline và logic hệ thống", 1)
    add_body(doc, "Pipeline dưới đây tách rõ nơi compressor tạo candidate và nơi ContextGuard quyết định candidate có được chấp nhận. Điểm bảo vệ quan trọng là deterministic guard chạy trước; semantic verifier chỉ được gọi khi kết quả deterministic là UNCERTAIN.")
    doc.add_picture(str(assets["pipeline"]), width=Inches(6.5))
    caption = doc.add_paragraph("Hình 1. ContextGuard pipeline: compressor không được tự promote candidate.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(size=9, color=MUTED)(caption.runs[0])
    add_heading(doc, "4.1 Logic quyết định", 2)
    add_table(
        doc,
        ["Bước", "Logic", "Failure-safe"],
        [
            ["Normalize/extract", "Chuẩn hóa text, fact, entity, relation, protected span", "Input invalid -> structured error"],
            ["Deterministic guards", "ExactGuard, LogicGuard, EntityGuard, RelationGuard", "Critical violation -> FAIL"],
            ["Policy mapping", "lenient/balanced/strict; strict mặc định", "FAIL/UNCERTAIN -> USE_ORIGINAL"],
            ["Semantic adapter", "Chỉ chạy khi deterministic UNCERTAIN", "Unavailable/error -> UNCERTAIN + fallback"],
            ["Report", "Status, violations, warnings, score, recommended action", "Không tự gán quality label"],
        ],
        [1900, 4800, 2660],
    )
    add_heading(doc, "4.2 API V1 và compatibility", 2)
    add_body(doc, "V1 routes vẫn là GET /v1/health, GET /v1/capabilities, POST /v1/analyze và POST /v1/validate. M10-M13 chỉ thêm scripts và artifact; không sửa breaking response field. Nếu cần thay đổi public contract, dùng /v2 thay vì âm thầm đổi V1.")

    add_heading(doc, "5. Môi trường và khả năng tái lập", 1)
    add_table(
        doc,
        ["Hạng mục", "Giá trị đã kiểm tra", "Tác động"],
        [
            ["OS / CPU / RAM", "Windows 11; i5-12500H; 24 GB", "Phù hợp native Windows + uv"],
            ["GPU", "RTX 3050 Laptop 4 GB", "CUDA venv ngoài repo; không coi là default runtime"],
            ["Python", "uv project 3.11.9; pyproject <3.12", "Global Python 3.12 không dùng cho benchmark"],
            ["Docker", "Docker Desktop 29.6.2; non-root image smoke pass", "R4 local only; chưa publish registry"],
            ["Data home", "D:\\fact_safeguard_data", "raw/extracted/normalized/cache/reviewer runs ngoài Git"],
        ],
        [1900, 4300, 3160],
    )
    add_body(doc, "Checksum, revision và license của nguồn M11 được ghi ở natural_adversarial_v1_manifest.json. Cleanup chỉ dọn .runtime/cache trong repo và không được xóa data home.")

    add_heading(doc, "6. M10a runtime benchmark", 1)
    add_body(doc, "M10a dùng 8 controlled cases: bốn bucket token short/medium/long/very_long × en/vi, đo cold/warm, process start, concurrency deterministic 1/2/4/8/16 và semantic 1/2/4. Manifest ghi rõ sáu semantic skips vì cặp input vượt giới hạn 512 token; không truncation âm thầm.")
    warm_ranges = {
        "R1 direct": "2.154-6,092.503 ms",
        "R2 ASGI": "3.260-318.427 ms",
        "R3 localhost": "154.538-4,088.412 ms",
        "R4 Docker": "151.062-3,854.673 ms",
        "R5 hybrid": "325.140-1,162.054 ms",
    }
    add_table(doc, ["Mode", "Warm P95 range", "Kết quả"], [[k, v, "error=0; timeout=0"] for k, v in warm_ranges.items()], [2500, 3000, 3860])
    add_body(doc, "Artifact: runtime_manifest.json, runtime_metrics.csv và runtime_samples.jsonl. API overhead được ghép theo cùng case/bucket/language/concurrency; số âm do contention/outlier không được diễn giải là API nhanh hơn direct.")
    add_callout(doc, "Giới hạn M10a", "R5 chỉ chạy hai short cases vì model pair max 512 token. Đây là skip có lý do, không phải full semantic coverage.")

    add_heading(doc, "7. M11 nguồn dữ liệu và benchmark 400 mẫu", 1)
    add_body(doc, "Pilot natural_adversarial_v1 được sinh ngoài Git với seed 20260729, exact 4 domain × 2 language × 2 construction label × 25 = 400. Split stratified là 280 dev và 120 hidden; hidden label tách khỏi payload reviewer.")
    add_table(
        doc,
        ["Nguồn", "License / revision", "Vai trò", "Trạng thái"],
        [
            ["ContractNLI", "CC-BY-4.0; stanfordnlp/contract-nli", "Business/legal clauses", "Downloaded + extracted"],
            ["QMSum", "MIT; Yale-LILY/QMSum", "Meeting summaries", "Downloaded"],
            ["QASPER", "CC-BY-4.0; HF pinned revision", "Scientific QA candidate", "Metadata-only; payload unavailable"],
            ["SciFact", "MIT; allenai/scifact", "Claim/evidence academic", "Downloaded + extracted"],
            ["XNLI", "MIT; facebookresearch/XNLI", "EN/VI NLI auxiliary", "Downloaded + extracted"],
            ["FastAPI + ContextGuard README", "MIT / public repo commit", "Technical documentation", "Captured + checksum"],
        ],
        [1900, 2300, 2900, 2260],
        font_size=8.6,
    )
    add_heading(doc, "7.1 Schema và label semantics", 2)
    add_body(doc, "Record có sample_id, domain, language, original_context, candidate_context, label/split, source_type, translated_marker, source_name, source_record_id, mutation_type, evaluation_role và notes. SAFE/UNSAFE là construction labels: SAFE giữ candidate source, UNSAFE dùng negation/conflict/instruction injection controlled. Đây chưa phải nhãn sự thật do human.")
    add_table(doc, ["Stratum", "Dev", "Hidden", "Ghi chú"], [["general × en/vi × SAFE/UNSAFE", "18 hoặc 17/cell", "phần còn lại", "balanced total"], ["academic × en/vi", "18 hoặc 17/cell", "phần còn lại", "SciFact/QMSum"], ["business × en/vi", "18 hoặc 17/cell", "phần còn lại", "ContractNLI/QMSum"], ["technical × en/vi", "18 hoặc 17/cell", "phần còn lại", "README docs"]], [3000, 1800, 1800, 2760])

    add_heading(doc, "8. M12 blind multi-agent AI review", 1)
    add_body(doc, "Năm vai trò độc lập nhận packet ngẫu nhiên: Academic, Technical, Business, VI-EN và Red-team. Mỗi packet có 400 dòng, chỉ gồm opaque item_id, domain, language, original và candidate. Reviewer không thấy label, mutation, source ID, prediction hoặc builder rationale.")
    add_table(doc, ["Metric", "Giá trị", "Cách đọc"], [["Unanimous rate", "0.575", "Năm reviewer cùng verdict"], ["Mean pairwise", "0.740", "Trung bình 10 cặp role"], ["Fleiss kappa", "0.521", "Agreement giữa AI reviewers"], ["Resolved majority", "347/400", "53 tie/unresolved bị loại"], ["Construction alignment", "0.859", "Không phải human truth validation"]], [2600, 1800, 4960])
    add_body(doc, "Raw JSONL reviewer runs nằm ngoài Git trong reviewer_runs. Public chỉ giữ m12_review_summary.json. Thuật ngữ đúng là AI-reviewed/AI-adjudicated; không claim human/domain expert.")

    add_heading(doc, "9. M13 compressor và token saving", 1)
    add_body(doc, "M13 chạy 24 record đại diện, RuleBasedCompressor và LLMLingua-2 ở rate .33/.50/.70, dùng protected spans và Qwen tokenizer. Bốn path đã đo: R6 direct compressor, R7 direct guard, R8 API guard và R9 direct+API. Tổng 672 rows; có ablation protected=true/false trên tám record đầu.")
    add_table(doc, ["Compressor", "Gross saving", "Safe saving", "PASS rate", "Fallback"], [["LLMLingua-2", "0.451", "0.024", "0.153", "24/672 (512-token guard)"], ["RuleBased", "0.150", "0.000", "0.417", "0"]], [2300, 1600, 1600, 1400, 2460])
    doc.add_picture(str(assets["safe_saving"]), width=Inches(6.5))
    caption = doc.add_paragraph("Hình 2. Safe effective token saving theo compressor/rate sau deterministic guard.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(size=9, color=MUTED)(caption.runs[0])
    add_callout(doc, "Cảnh báo diễn giải", "Gross saving cao không đồng nghĩa candidate an toàn. Safe saving chỉ được tính khi guard trả PASS; fallback không được ghi nhận là success. Break-even downstream chưa đo vì chưa có target LLM/API cost contract.")

    add_heading(doc, "10. Demo context trước và sau", 1)
    demo = next(row for row in dev if row["domain"] == "technical" and row["language"] == "en")
    guard = ContextGuard(GuardConfig(language="en", profile="technical", policy="strict"))
    spans = guard.analyze(demo["original_context"]).protected_spans
    compressed = RuleBasedCompressor().compress(demo["original_context"], spans)
    add_table(doc, ["Stage", "Context excerpt"], [["Original", demo["original_context"][:950]], ["RuleBased candidate", compressed[:950]], ["Guard", "Measured through R7/R8; status is recorded per sample, not inferred from construction label"]], [1900, 7460], font_size=8.7)
    add_body(doc, "Protected spans ở demo được lấy từ analyzer, không hard-code vào report. Với M13, original/candidate/token count/status/fallback được lưu per sample trong m13_compressor_samples.jsonl.")

    add_heading(doc, "11. Demo mẫu từng dataset", 1)
    add_body(doc, "Các excerpt dưới đây là minh họa nhỏ để đọc pipeline. Chúng không thay thế toàn bộ dataset và không phải quality claim.")
    source_rows = []
    for source in ("XNLI", "SciFact", "ContractNLI", "QMSum", "technical_documentation"):
        row = next((item for item in dev if item["source_name"] == source), None)
        if row:
            source_rows.append([source, row["domain"], row["language"], row["original_context"][:230], row["candidate_context"][:230]])
        else:
            source_rows.append([source, "-", "-", "No public payload sample in pilot", "QASPER is metadata-only at pinned revision"])
    add_table(doc, ["Dataset", "Domain", "Lang", "Original excerpt", "Candidate excerpt"], source_rows, [1600, 1100, 700, 2980, 2980], font_size=7.7)
    add_body(doc, "QASPER được cố ý hiển thị là metadata-only, không bịa sample. Với VI records từ English source, translated_marker=true và cần translation QA trước khi gọi là naturally Vietnamese.")

    add_heading(doc, "12. Review chéo và rủi ro còn lại", 1)
    add_table(doc, ["Rủi ro / xung đột", "Cách xử lý", "Còn lại"], [["M10 R6-R9 overlap", "Khóa scope: R1-R5 ở M10a, R6-R9 ở M13", "Không còn overlap logic"], ["AI review bị gọi nhầm human", "Đổi thuật ngữ thành blind multi-agent AI review", "Không có human annotation"], ["QASPER thiếu payload", "Manifest metadata-only + limitation", "Cần external access/revision mới"], ["Model input >512", "Explicit LLMLingua fallback reason", "24/672 fallback"], ["Raw data public leak", "Data home ngoài repo; public manifest sanitize path", "M14 scan tiếp"]], [2700, 3900, 2760])
    add_heading(doc, "12.1 Các claim không được phép", 2)
    for item in (
        "Không gọi SAFE/UNSAFE construction label là human truth hoặc gold label production.",
        "Không gọi M12 là human expert review; chỉ là AI reviewer agreement.",
        "Không suy ra cost saving hoặc break-even khi chưa có downstream target LLM contract.",
        "Không suy ra semantic generalization từ 2 short cases M10a hoặc controlled pilot M13.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "13. M14 audit và điều kiện kết thúc", 1)
    add_table(doc, ["Gate", "Điều kiện", "Artifact / lệnh"], [["Test", "pytest, Ruff, mypy, uv lock/pip check", "uv run pytest; uv run ruff check ."], ["Security", "pip-audit, secret/large/forbidden/model scan", "M14 audit script"], ["Data home", "raw/extracted/normalized/cache/reviewer_runs tồn tại", "D:\\fact_safeguard_data"], ["Docs", "CONTEXT/README đồng bộ với artifact thực", "exactly 3 Markdown"], ["Git", "clean, commits rõ, push master", "git status; git log; git push"], ["Word", "render PNG mọi trang + inspect", "render_docx.py"]], [1900, 4300, 3160])
    add_body(doc, "Status được chọn trong enum: COMPLETED, COMPLETED_WITH_LIMITATIONS, PARTIALLY_COMPLETED hoặc BLOCKED_BY_EXTERNAL_ACCESS. Evidence M14 phù hợp với COMPLETED_WITH_LIMITATIONS vì QASPER payload, human annotation, downstream break-even và natural-language coverage rộng vẫn chưa có.")

    add_heading(doc, "14. Artifact catalog public", 1)
    public_files = sorted(str(path.relative_to(ROOT)) for path in (ROOT / "artifacts/final").glob("*"))
    add_table(doc, ["Artifact", "Vai trò"], [[path, "Public evidence; không chứa raw/model cache" if path.endswith((".json", ".csv", ".jsonl")) else "Public report input"] for path in public_files], [4300, 5060], font_size=8.7)
    add_body(doc, "Các file raw/normalized/reviewer runs không nằm trong artifact catalog vì được giữ tại data home ngoài repository. Report chỉ nhúng aggregate và excerpt đã giới hạn độ dài.")

    add_heading(doc, "15. Nguồn tham khảo", 1)
    refs = [
        "ContractNLI official repository: https://github.com/stanfordnlp/contract-nli (CC-BY-4.0).",
        "QMSum official repository: https://github.com/Yale-LILY/QMSum (MIT; meeting summarization).",
        "SciFact official repository: https://github.com/allenai/scifact (MIT; claim/evidence).",
        "XNLI official repository: https://github.com/facebookresearch/XNLI (MIT; multilingual NLI incl. Vietnamese).",
        "QASPER pinned dataset page: https://huggingface.co/datasets/allenai/qasper/tree/5475f6d48704155776600dcf183a3ddb05800539/qasper (CC-BY-4.0; payload availability limitation recorded).",
        "FastAPI official repository: https://github.com/fastapi/fastapi.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    add_heading(doc, "Phụ lục A. Lệnh tái lập", 1)
    for command in (
        "uv run python scripts/runtime_benchmark.py --output .runtime/runtime_m10_final --local-files-only",
        "uv run python scripts/dataset_pipeline.py --validate",
        "uv run python scripts/prepare_review.py",
        "uv run python scripts/aggregate_review.py",
        "$env:HF_HOME='D:\\AI_Cache\\huggingface'; uv run python scripts/compressor_benchmark.py --output .runtime/m13_compressor --limit 24 --local-files-only",
        "uv run pytest -q; uv run ruff check .; uv run mypy src",
    ):
        paragraph = doc.add_paragraph(style="No Spacing")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(command)
        _font(name="Consolas", size=9, color="344054")(run)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    print(build())
